"""
导出路由：HTML → Word (.docx) / PDF
- Word：用 python-docx 递归解析 HTML，完整保留格式（嵌套粗斜体、表格、图片、列表等）
- PDF：WeasyPrint 直接渲染 HTML，质量最佳
"""
import base64
import io
from typing import Optional
from urllib.parse import quote

# WeasyPrint 在 Windows 需要 GTK 运行库，懒加载避免未安装时阻止后端启动
try:
    import weasyprint as _weasyprint
    _WEASYPRINT_OK = True
except OSError:
    _weasyprint = None  # type: ignore
    _WEASYPRINT_OK = False

from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from fastapi import APIRouter
from fastapi.responses import StreamingResponse
from pydantic import BaseModel

router = APIRouter()

# ─── 请求模型 ──────────────────────────────────────────────────────────────────


class ExportRequest(BaseModel):
    project_name: str
    html: str  # exportToHtml() 生成的完整 HTML 字符串


# ─── HTML → DOCX ───────────────────────────────────────────────────────────────


def _add_hr(doc: Document):
    """在段落底部加横线模拟 <hr>。"""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "auto")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _process_inline(
    node: Tag,
    para,
    bold: bool = False,
    italic: bool = False,
    underline: bool = False,
    strike: bool = False,
    code: bool = False,
):
    """递归处理行内节点，正确继承嵌套格式（如 <strong><em>文字</em></strong>）。"""
    for child in node.children:
        if isinstance(child, NavigableString):
            text = str(child)
            if not text:
                continue
            run = para.add_run(text)
            run.bold = bold or None
            run.italic = italic or None
            run.underline = underline or None
            run.font.strike = strike or None
            if code:
                run.font.name = "Courier New"
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
        elif isinstance(child, Tag):
            t = child.name.lower()
            _process_inline(
                child,
                para,
                bold=bold or t in ("strong", "b"),
                italic=italic or t in ("em", "i"),
                underline=underline or t == "u",
                strike=strike or t in ("s", "strike", "del"),
                code=code or t == "code",
            )


def _get_align(style: str) -> Optional[WD_ALIGN_PARAGRAPH]:
    if "text-align: center" in style or "text-align:center" in style:
        return WD_ALIGN_PARAGRAPH.CENTER
    if "text-align: right" in style or "text-align:right" in style:
        return WD_ALIGN_PARAGRAPH.RIGHT
    if "text-align: justify" in style or "text-align:justify" in style:
        return WD_ALIGN_PARAGRAPH.JUSTIFY
    return None


def _process_block(node: Tag, doc: Document):
    """递归将块级 HTML 节点转换为 Word 段落。"""
    if isinstance(node, NavigableString):
        text = str(node).strip()
        if text:
            p = doc.add_paragraph()
            p.add_run(text)
        return

    if not isinstance(node, Tag):
        return

    tag = node.name.lower()

    # 标题
    if tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
        level = int(tag[1])
        p = doc.add_heading("", level=min(level, 9))
        _process_inline(node, p)
        return

    # 段落
    if tag == "p":
        p = doc.add_paragraph()
        align = _get_align(node.get("style", ""))
        if align:
            p.alignment = align
        _process_inline(node, p)
        return

    # 无序列表
    if tag == "ul":
        for li in node.find_all("li", recursive=False):
            p = doc.add_paragraph(style="List Bullet")
            _process_inline(li, p)
        return

    # 有序列表
    if tag == "ol":
        for li in node.find_all("li", recursive=False):
            p = doc.add_paragraph(style="List Number")
            _process_inline(li, p)
        return

    # 引用块
    if tag == "blockquote":
        p = doc.add_paragraph(style="Quote")
        _process_inline(node, p)
        return

    # 代码块
    if tag == "pre":
        code_text = node.get_text("\n")
        p = doc.add_paragraph(style="No Spacing")
        run = p.add_run(code_text)
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        p.paragraph_format.left_indent = Cm(1)
        # 灰色背景用段落底纹
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "F5F5F5")
        pPr.append(shd)
        return

    # 表格
    if tag == "table":
        rows = node.find_all("tr")
        if not rows:
            return
        col_count = max(len(r.find_all(["td", "th"])) for r in rows)
        if col_count == 0:
            return
        tbl = doc.add_table(rows=len(rows), cols=col_count)
        tbl.style = "Table Grid"
        for i, tr in enumerate(rows):
            cells = tr.find_all(["td", "th"])
            for j, cell in enumerate(cells):
                if j < col_count:
                    tc = tbl.rows[i].cells[j]
                    tc.text = ""
                    p = tc.paragraphs[0]
                    _process_inline(cell, p, bold=cell.name == "th")
        return

    # 图片
    if tag == "img":
        src = node.get("src", "")
        if src.startswith("data:image"):
            try:
                _, data = src.split(",", 1)
                img_bytes = base64.b64decode(data)
                doc.add_picture(io.BytesIO(img_bytes), width=Inches(5.5))
            except Exception:
                pass
        return

    # 分割线
    if tag == "hr":
        _add_hr(doc)
        return

    # div / section 等容器 → 递归子节点
    for child in node.children:
        _process_block(child, doc)


def html_to_docx(project_name: str, html: str) -> bytes:
    doc = Document()

    # 页面设置（A4）
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)

    # 书名作为封面标题
    title_p = doc.add_heading(project_name, level=0)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    soup = BeautifulSoup(html, "lxml")
    body = soup.find("body") or soup

    # 跳过 <h1>（即书名，已在上方插入）
    for child in body.children:
        if isinstance(child, Tag) and child.name == "h1":
            continue
        _process_block(child, doc)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ─── 路由 ──────────────────────────────────────────────────────────────────────


@router.post("/docx")
async def export_docx(req: ExportRequest):
    docx_bytes = html_to_docx(req.project_name, req.html)
    filename = quote(f"{req.project_name}.docx")
    return StreamingResponse(
        io.BytesIO(docx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{filename}"},
    )


@router.get("/pdf-status")
def pdf_status():
    """前端可用此接口检测 PDF 功能是否可用。"""
    return {"available": _WEASYPRINT_OK}


@router.post("/pdf")
async def export_pdf(req: ExportRequest):
    if not _WEASYPRINT_OK:
        from fastapi import HTTPException
        raise HTTPException(
            status_code=503,
            detail=(
                "PDF 导出需要 WeasyPrint 及 GTK 运行库。\n"
                "Windows 安装步骤：\n"
                "1. 下载 GTK3 Runtime：https://github.com/tschoonj/GTK-for-Windows-Runtime-Environment-Installer/releases\n"
                "2. 安装后重启后端即可。\n"
                "详见：https://doc.courtbouillon.org/weasyprint/stable/first_steps.html"
            ),
        )
    pdf_bytes = _weasyprint.HTML(string=req.html).write_pdf()
    filename = quote(f"{req.project_name}.pdf")
    return StreamingResponse(
        io.BytesIO(pdf_bytes),
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{filename}"},
    )
