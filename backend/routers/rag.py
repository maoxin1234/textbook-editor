"""
RAG 模块：
1. 上传参考文档（PDF / DOCX / TXT / MD）
2. 提取文本 → 分块 → 本地 ONNX 嵌入 → 存入 ChromaDB
3. 检索接口供 AI 对话使用
"""
import io
import re
import uuid

import chromadb
from chromadb.utils.embedding_functions import DefaultEmbeddingFunction
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select

from database import get_db
from models import User, Project, RagDocument
from schemas import RagDocumentOut
from deps import get_current_user

router = APIRouter()

CHROMA_PATH = "./chroma_db"
MAX_FILE_SIZE = 10 * 1024 * 1024  # 10 MB
CHUNK_SIZE = 500       # 每块目标字符数
CHUNK_OVERLAP = 50     # 块间重叠字符数

ALLOWED_TYPES = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "text/plain",
    "text/markdown",
    "text/x-markdown",
}

# ─── ChromaDB 工厂 ────────────────────────────────────────────────────────────

def _get_collection(project_id: str):
    """每个项目独享一个 ChromaDB collection，使用本地 ONNX 嵌入模型（首次用时自动下载 ~90MB）。"""
    client = chromadb.PersistentClient(path=CHROMA_PATH)
    return client.get_or_create_collection(
        name=f"proj_{project_id.replace('-', '_')}",
        embedding_function=DefaultEmbeddingFunction(),
        metadata={"hnsw:space": "cosine"},
    )

# ─── 文本提取 ─────────────────────────────────────────────────────────────────

def _extract_text(content: bytes, content_type: str) -> str:
    if content_type == "application/pdf":
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(content))
        return "\n".join(
            page.extract_text() or "" for page in reader.pages
        )
    if content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        from docx import Document
        doc = Document(io.BytesIO(content))
        return "\n".join(p.text for p in doc.paragraphs)
    # TXT / MD
    return content.decode("utf-8", errors="ignore")

# ─── 文本分块（中英文通用） ───────────────────────────────────────────────────

def _split_text(text: str) -> list[str]:
    # 先按段落切分
    paragraphs = [p.strip() for p in re.split(r"\n{2,}", text) if p.strip()]
    chunks: list[str] = []
    buf = ""

    for para in paragraphs:
        if len(buf) + len(para) <= CHUNK_SIZE:
            buf = (buf + "\n" + para).strip()
        else:
            if buf:
                chunks.append(buf)
            # 段落本身超长：按句子进一步切分
            if len(para) > CHUNK_SIZE:
                sentences = re.split(r"(?<=[。！？.!?])", para)
                for sent in sentences:
                    if len(buf) + len(sent) <= CHUNK_SIZE:
                        buf = (buf + sent).strip()
                    else:
                        if buf:
                            chunks.append(buf)
                        buf = sent.strip()
            else:
                buf = para

    if buf:
        chunks.append(buf)

    # 加重叠（每块结尾附上下一块开头的 CHUNK_OVERLAP 字符）
    overlapped = []
    for i, chunk in enumerate(chunks):
        if i < len(chunks) - 1:
            overlap = chunks[i + 1][:CHUNK_OVERLAP]
            overlapped.append(chunk + overlap)
        else:
            overlapped.append(chunk)

    return [c for c in overlapped if c.strip()]

# ─── 路由 ─────────────────────────────────────────────────────────────────────

@router.get("/{project_id}/documents", response_model=list[RagDocumentOut])
async def list_documents(
    project_id: str,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    await _assert_owner(project_id, current_user.id, db)
    result = await db.execute(
        select(RagDocument)
        .where(RagDocument.project_id == project_id)
        .order_by(RagDocument.created_at.desc())
    )
    return result.scalars().all()


@router.post("/{project_id}/documents", response_model=RagDocumentOut, status_code=201)
async def upload_document(
    project_id: str,
    file: UploadFile = File(...),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    await _assert_owner(project_id, current_user.id, db)

    content_type = file.content_type or "text/plain"
    if content_type not in ALLOWED_TYPES:
        raise HTTPException(status_code=415, detail=f"不支持的文件类型：{content_type}，支持 PDF / DOCX / TXT / MD")

    content = await file.read()
    if len(content) > MAX_FILE_SIZE:
        raise HTTPException(status_code=413, detail="文件超过 10MB 限制")

    # 提取文本 → 分块
    text = _extract_text(content, content_type)
    chunks = _split_text(text)
    if not chunks:
        raise HTTPException(status_code=422, detail="文档内容为空，无法处理")

    # 写入 ChromaDB
    doc_id = str(uuid.uuid4())
    collection = _get_collection(project_id)
    collection.add(
        documents=chunks,
        ids=[f"{doc_id}_{i}" for i in range(len(chunks))],
        metadatas=[
            {"doc_id": doc_id, "filename": file.filename or "unknown", "chunk_index": i}
            for i in range(len(chunks))
        ],
    )

    # 写入数据库
    rag_doc = RagDocument(
        id=doc_id,
        project_id=project_id,
        filename=file.filename or "unknown",
        content_type=content_type,
        chunk_count=len(chunks),
    )
    db.add(rag_doc)
    await db.commit()
    await db.refresh(rag_doc)
    return rag_doc


@router.delete("/{project_id}/documents/{doc_id}", status_code=204)
async def delete_document(
    project_id: str,
    doc_id: str,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    await _assert_owner(project_id, current_user.id, db)
    rag_doc = await db.get(RagDocument, doc_id)
    if not rag_doc or rag_doc.project_id != project_id:
        raise HTTPException(status_code=404, detail="文档不存在")

    # 从 ChromaDB 删除该文档的所有 chunks
    try:
        collection = _get_collection(project_id)
        results = collection.get(where={"doc_id": doc_id})
        if results["ids"]:
            collection.delete(ids=results["ids"])
    except Exception:
        pass

    await db.delete(rag_doc)
    await db.commit()


# ─── 供 AI 路由调用的检索函数 ─────────────────────────────────────────────────

def retrieve_context(project_id: str, query: str, top_k: int = 4) -> str:
    """检索与 query 最相关的 top_k 块，拼接为上下文字符串。"""
    try:
        collection = _get_collection(project_id)
        if collection.count() == 0:
            return ""
        results = collection.query(query_texts=[query], n_results=min(top_k, collection.count()))
        docs = results.get("documents", [[]])[0]
        if not docs:
            return ""
        return "\n\n---\n\n".join(docs)
    except Exception:
        return ""


# ─── 工具 ─────────────────────────────────────────────────────────────────────

async def _assert_owner(project_id: str, user_id: int, db: AsyncSession):
    project = await db.get(Project, project_id)
    if not project or project.user_id != user_id:
        raise HTTPException(status_code=404, detail="项目不存在")
